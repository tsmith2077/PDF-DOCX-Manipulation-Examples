# Creates docx file of custom invitations for 
# docx file list of names.

import docx, os
from docx.shared import Pt


def createInvitations(filename):
    guestsDoc = open(os.path.abspath(filename))
    guestList = guestsDoc.readlines()
    invitationsDoc = docx.Document()
    for name in range(len(guestList)):
        guestName = guestList[name].strip('\n')
        invitationsDoc.add_paragraph('It would be a pleasure to have the company of\n')
        invitationsDoc.paragraphs[name].runs[0].italic = True
        invitationsDoc.paragraphs[name].add_run(guestName).bold = True
        invitationsDoc.paragraphs[name].add_run('\nat 11010 Memory Lane on the Evening of\n').italic = True
        invitationsDoc.paragraphs[name].add_run('April 1st\n').style = None
        invitationsDoc.paragraphs[name].add_run("at 7 o'clock").italic = True
        invitationsDoc.paragraphs[name].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        invitationsDoc.paragraphs[name].alignment = 1
        invitationsDoc.paragraphs[name].style.font.size = Pt(18)
    invitationsDoc.save('invitations.docx')
        
        
        